"""
제안서 품질 관리 페이지
"""
import streamlit as st
import json
import time
import io
from datetime import datetime
from docx import Document
import pandas as pd
import PyPDF2
import pdfplumber

def show():
    """제안서 품질 관리 페이지 표시"""
    st.title("📋 제안서 품질 관리")
    st.markdown("제안서의 품질을 검증하고 개선 방안을 제시합니다.")
    
    # 분석 모드 선택
    analysis_mode = st.radio(
        "분석 모드 선택",
        ["저장된 RFP 기반 품질 검증", "직접 입력 기반 품질 검증"],
        horizontal=True
    )
    
    if analysis_mode == "저장된 RFP 기반 품질 검증":
        show_stored_rfp_quality_check()
    else:
        show_manual_input_quality_check()

def show_stored_rfp_quality_check():
    """저장된 RFP 기반 품질 검증"""
    st.subheader("📁 저장된 RFP 기반 품질 검증")
    
    try:
        azure_services = st.session_state.azure_services
        directories = azure_services.get_directories()
        
        if not directories:
            st.warning("저장된 RFP가 없습니다.")
            return
        
        # 디렉토리 목록을 테이블로 표시
        st.subheader("📋 저장된 RFP 목록")
        
        # 페이지네이션 설정
        items_per_page = 10
        total_pages = (len(directories) + items_per_page - 1) // items_per_page
        
        if total_pages > 1:
            page = st.selectbox("페이지 선택", range(1, total_pages + 1), key="quality_page")
        else:
            page = 1
        
        start_idx = (page - 1) * items_per_page
        end_idx = start_idx + items_per_page
        current_directories = directories[start_idx:end_idx]
        
        # 테이블 데이터 생성 (선택 컬럼 포함)
        table_data = []
        for i, directory in enumerate(current_directories):
            table_data.append({
                "선택": f"RFP {i+1}",
                "한글명": directory['korean_name'],
                "생성일": directory['created_date'],
                "프로젝트 요약": directory['project_summary'][:50] + "..." if len(directory['project_summary']) > 50 else directory['project_summary']
            })
        
        # 테이블과 라디오 버튼을 함께 표시
        st.markdown("**📋 저장된 RFP 목록에서 검증할 RFP를 선택하세요:**")
        
        # 컨테이너로 테이블과 라디오 버튼을 그룹화
        with st.container():
            # 테이블 표시
            df = pd.DataFrame(table_data)
            st.dataframe(df, width='stretch')
            
            # 라디오 버튼을 테이블 바로 아래에 배치
            selected_index = st.radio(
                "검증할 RFP 선택:",
                range(len(current_directories)),
                format_func=lambda x: f"🔍 {current_directories[x]['korean_name']} ({current_directories[x]['created_date']})",
                key="quality_rfp_selection"
            )
        
        if selected_index is not None:
            selected_directory = current_directories[selected_index]
            
            # 현재 작업 디렉토리를 세션 상태에 저장
            st.session_state.current_directory = selected_directory['name']
            st.session_state.current_container = "rfp-documents"
            
            # 제안서 파일 업로드
            st.subheader("제안서 파일 업로드")
            uploaded_proposal = st.file_uploader(
                "제안서 파일 업로드",
                type=['pdf', 'docx', 'txt'],
                help="검증할 제안서 파일을 업로드하세요."
            )
            
            if uploaded_proposal and st.button("품질 검증 시작", type="primary"):
                with st.spinner("제안서 품질을 검증하고 있습니다..."):
                    run_quality_check(selected_directory['name'], uploaded_proposal)
                
                st.success("품질 검증이 완료되었습니다!")
                
    except Exception as e:
        st.error(f"오류가 발생했습니다: {str(e)}")

def show_manual_input_quality_check():
    """직접 입력 기반 품질 검증"""
    st.subheader("📝 직접 입력 기반 품질 검증")
    
    # RFP 파일 업로드
    st.subheader("RFP 파일 업로드")
    uploaded_rfp = st.file_uploader(
        "RFP 문서 업로드",
        type=['pdf', 'docx', 'txt'],
        help="검증 기준이 될 RFP 문서를 업로드하세요.",
        key="rfp_upload"
    )
    
    if uploaded_rfp:
        st.info(f"📁 업로드된 RFP: {uploaded_rfp.name} ({uploaded_rfp.size} bytes)")
        
        # 제안서 파일 업로드
        st.subheader("제안서 파일 업로드")
        uploaded_proposal = st.file_uploader(
            "제안서 파일 업로드",
            type=['pdf', 'docx', 'txt'],
            help="검증할 제안서 파일을 업로드하세요.",
            key="proposal_upload"
        )
        
        if uploaded_proposal:
            st.info(f"📁 업로드된 제안서: {uploaded_proposal.name} ({uploaded_proposal.size} bytes)")
            
            # 추가 입력 폼
            col1, col2 = st.columns(2)
            
            with col1:
                industry = st.selectbox("업종", ["은행", "보험", "증권", "카드"])
                business_characteristics = st.text_area(
                    "비즈니스 특성 (선택사항)",
                    placeholder="예: 대형 은행, 디지털 전환 추진, 고객 경험 개선 중점"
                )
            
            with col2:
                rfp_summary = st.text_area(
                    "RFP 요약 (선택사항)",
                    placeholder="RFP의 주요 내용을 간단히 요약해주세요. (비워두면 자동 분석됩니다)",
                    height=100
                )
            
            if st.button("품질 검증 시작", type="primary"):
                with st.spinner("제안서 품질을 검증하고 있습니다..."):
                    run_manual_quality_check_with_files(uploaded_rfp, uploaded_proposal, industry, business_characteristics, rfp_summary)
                
                st.success("품질 검증이 완료되었습니다!")

def run_manual_quality_check_with_files(uploaded_rfp, uploaded_proposal, industry, business_characteristics, rfp_summary):
    """파일 업로드 기반 품질 검증 실행"""
    try:
        # RFP 내용 추출
        rfp_content = uploaded_rfp.read()
        if uploaded_rfp.type == "text/plain":
            rfp_text = rfp_content.decode('utf-8', errors='ignore')
        else:
            # 파일 확장자에 따라 적절한 텍스트 추출 방법 선택
            if uploaded_rfp.name.lower().endswith('.pdf'):
                rfp_text = extract_text_from_pdf_bytes(rfp_content)
            elif uploaded_rfp.name.lower().endswith('.docx'):
                rfp_text = extract_text_from_docx_bytes(rfp_content)
            else:
                rfp_text = f"[{uploaded_rfp.name} 파일 내용 - {len(rfp_content)} bytes]"
        
        # 제안서 내용 추출
        proposal_content = uploaded_proposal.read()
        if uploaded_proposal.type == "text/plain":
            proposal_text = proposal_content.decode('utf-8', errors='ignore')
        else:
            # 파일 확장자에 따라 적절한 텍스트 추출 방법 선택
            if uploaded_proposal.name.lower().endswith('.pdf'):
                proposal_text = extract_text_from_pdf_bytes(proposal_content)
            elif uploaded_proposal.name.lower().endswith('.docx'):
                proposal_text = extract_text_from_docx_bytes(proposal_content)
            else:
                proposal_text = f"[{uploaded_proposal.name} 파일 내용 - {len(proposal_content)} bytes]"
        
        # RFP 요약이 없으면 자동 생성
        if not rfp_summary:
            rfp_summary = generate_auto_rfp_summary_for_quality(rfp_text, industry)
        
        rfp_info = {
            "industry": industry,
            "business_characteristics": business_characteristics or "일반적인 금융 서비스",
            "rfp_summary": rfp_summary,
            "rfp_content": rfp_text
        }
        
        # 품질 검증 실행 (수동 입력 방식)
        generate_quality_results_manual(rfp_info, proposal_text)
        
    except Exception as e:
        st.error(f"품질 검증 중 오류: {str(e)}")

def generate_auto_rfp_summary_for_quality(content, industry):
    """품질 검증용 RFP 자동 요약 생성"""
    try:
        azure_services = st.session_state.azure_services
        
        messages = [
            {
                "role": "system",
                "content": f"당신은 {industry} 업종의 RFP 분석 전문가입니다. 품질 검증을 위한 RFP 요약을 생성해주세요."
            },
            {
                "role": "user",
                "content": f"""
                업종: {industry}
                RFP 내용:
                {content}
                
                다음 형식으로 RFP를 요약해주세요:
                
                ## RFP 요약
                - 프로젝트 목표: [핵심 목표]
                - 주요 요구사항: [주요 요구사항 3-5개]
                - 기술 스택: [주요 기술 스택]
                - 예상 규모: [소규모/중규모/대규모]
                """
            }
        ]
        
        return azure_services.call_openai(messages)
    except Exception as e:
        return f"RFP 자동 요약 생성 중 오류: {str(e)}"

def run_quality_check(directory_name, uploaded_proposal):
    """RFP 기반 품질 검증 실행"""
    try:
        azure_services = st.session_state.azure_services
        
        # 제안서를 main_proposal_ 형식으로 업로드
        proposal_content_bytes = uploaded_proposal.read()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        proposal_filename = f"main_proposal_{timestamp}_{uploaded_proposal.name}"
        
        # 선택한 디렉토리에 제안서 업로드
        container_name = "rfp-documents"
        azure_services.upload_file_to_directory(container_name, directory_name, proposal_filename, proposal_content_bytes)
        st.success(f"제안서가 업로드되었습니다: {proposal_filename}")
        
        # RFP 정보 추출
        rfp_info = extract_rfp_info_from_directory(directory_name)
        
        # 품질 검증 실행 (URL 기반)
        generate_quality_results_with_urls(rfp_info, directory_name)
        
    except Exception as e:
        st.error(f"품질 검증 중 오류: {str(e)}")

def extract_rfp_info_from_directory(directory_name):
    """디렉토리에서 RFP 정보 추출"""
    try:
        azure_services = st.session_state.azure_services
        container_name = "rfp-documents"
        metadata = azure_services.get_directory_metadata_from_path(container_name, directory_name)
        
        return {
            "name": metadata.get('korean_name', directory_name),
            "directory_name": directory_name,
            "date": metadata.get('created_date', ''),
            "industry": "은행",  # 기본값
            "project_summary": metadata.get('project_summary', ''),
            "metadata": metadata
        }
    except Exception as e:
        return {
            "name": directory_name,
            "directory_name": directory_name,
            "date": "",
            "industry": "은행",
            "project_summary": ""
        }


def generate_quality_results_manual(rfp_info, proposal_content):
    """수동 입력 기반 품질 검증 결과 생성"""
    try:
        # 탭으로 결과 표시
        tab1, tab2 = st.tabs([
            "RFP 요구사항 vs 제안서 매핑", 
            "누락 항목 자동 감지"
        ])
        
        # 병렬 처리를 위한 placeholder 생성
        with tab1:
            mapping_placeholder = st.empty()
            mapping_placeholder.info("🔄 요구사항 매핑 분석 중...")
        
        with tab2:
            missing_placeholder = st.empty()
            missing_placeholder.info("🔄 누락 항목 감지 중...")
        
        # Azure 서비스를 미리 가져와서 병렬 처리에서 사용할 수 있도록 준비
        azure_services = st.session_state.azure_services
        
        # 병렬 처리로 모든 분석 동시 실행
        import concurrent.futures
        
        def run_manual_quality_analysis(azure_services, rfp_info, proposal_content):
            """수동 입력 기반 품질 분석 실행"""
            with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
                # 모든 분석을 병렬로 실행
                future_mapping = executor.submit(generate_requirements_mapping_manual, azure_services, rfp_info, proposal_content)
                future_missing = executor.submit(detect_missing_items_manual, azure_services, rfp_info, proposal_content)
                
                # 결과를 순차적으로 표시 (완료되는 대로)
                results = {}
                
                # 요구사항 매핑 결과 표시
                mapping_result = future_mapping.result()
                results['mapping_result'] = mapping_result
                mapping_placeholder.markdown(
                    f"""
                    <div style="max-height: 600px; overflow-y: auto; padding: 15px; border: 1px solid #e0e0e0; border-radius: 8px; background-color: #fafafa;">
                        {mapping_result}
                    </div>
                    """, 
                    unsafe_allow_html=True
                )
                
                # 누락 항목 결과 표시
                missing_items = future_missing.result()
                results['missing_items'] = missing_items
                missing_placeholder.markdown(
                    f"""
                    <div style="max-height: 600px; overflow-y: auto; padding: 15px; border: 1px solid #e0e0e0; border-radius: 8px; background-color: #fafafa;">
                        {missing_items}
                    </div>
                    """, 
                    unsafe_allow_html=True
                )
                
                return results
        
        # 분석 실행
        results = run_manual_quality_analysis(azure_services, rfp_info, proposal_content)
        mapping_result = results['mapping_result']
        missing_items = results['missing_items']
        
        # 다운로드 버튼을 탭 밖으로 이동
        st.subheader("📥 품질 검증 결과 다운로드")
        
        # 통합 품질 검증 결과 생성
        combined_content = f"""
# 제안서 품질 검증 결과

## 1. 요구사항 매핑 결과
{mapping_result}

## 2. 누락 항목 분석
{missing_items}

        """
        
        # 다운로드 데이터 미리 생성
        quality_download_data = create_quality_download_data("rfp_quality_check.docx", combined_content)
        
        # HTML 다운로드 링크 사용 (페이지 리로드 방지)
        quality_link = create_quality_download_link(quality_download_data, "rfp_quality_check.docx", "📋 품질 검증 결과 다운로드")
        st.markdown(quality_link, unsafe_allow_html=True)
                
    except Exception as e:
        st.error(f"품질 검증 결과 생성 중 오류: {str(e)}")

def generate_quality_results_with_urls(rfp_info, directory_name):
    """URL 기반 품질 검증 결과 생성"""
    try:
        azure_services = st.session_state.azure_services
        container_name = "rfp-documents"
        
        # 파일 URL 생성
        main_rfp_url = f"{container_name}/{directory_name}/main_rfp_"
        main_proposal_url = f"{container_name}/{directory_name}/main_proposal_"
        
        # 탭으로 결과 표시
        tab1, tab2 = st.tabs([
            "RFP 요구사항 vs 제안서 매핑", 
            "누락 항목 자동 감지"
        ])
        
        # 병렬 처리를 위한 placeholder 생성
        with tab1:
            mapping_placeholder = st.empty()
            mapping_placeholder.info("🔄 요구사항 매핑 분석 중...")
        
        with tab2:
            missing_placeholder = st.empty()
            missing_placeholder.info("🔄 누락 항목 감지 중...")
        
        # 병렬 처리로 모든 분석 동시 실행
        import concurrent.futures
        
        def run_quality_analysis_with_urls(azure_services, rfp_info, main_rfp_url, main_proposal_url):
            """URL 기반 품질 분석 실행"""
            with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
                # 모든 분석을 병렬로 실행
                future_mapping = executor.submit(generate_requirements_mapping_with_urls, azure_services, rfp_info, main_rfp_url, main_proposal_url)
                future_missing = executor.submit(detect_missing_items_with_urls, azure_services, rfp_info, main_rfp_url, main_proposal_url)
                
                # 결과를 순차적으로 표시 (완료되는 대로)
                results = {}
                
                # 요구사항 매핑 결과 표시
                mapping_result = future_mapping.result()
                results['mapping_result'] = mapping_result
                mapping_placeholder.markdown(
                    f"""
                    <div style="max-height: 600px; overflow-y: auto; padding: 15px; border: 1px solid #e0e0e0; border-radius: 8px; background-color: #fafafa;">
                        {mapping_result}
                    </div>
                    """, 
                    unsafe_allow_html=True
                )
                
                # 누락 항목 결과 표시
                missing_items = future_missing.result()
                results['missing_items'] = missing_items
                missing_placeholder.markdown(
                    f"""
                    <div style="max-height: 600px; overflow-y: auto; padding: 15px; border: 1px solid #e0e0e0; border-radius: 8px; background-color: #fafafa;">
                        {missing_items}
                    </div>
                    """, 
                    unsafe_allow_html=True
                )
                
                return results
        
        # 분석 실행
        results = run_quality_analysis_with_urls(azure_services, rfp_info, main_rfp_url, main_proposal_url)
        mapping_result = results['mapping_result']
        missing_items = results['missing_items']
        
        # 품질 검증 결과를 디렉토리에 자동 저장
        save_quality_results_to_directory(mapping_result, missing_items)
        
        # 다운로드 버튼을 탭 밖으로 이동
        st.subheader("📥 품질 검증 결과 다운로드")
        
        # 통합 품질 검증 결과 생성
        combined_content = f"""
# 제안서 품질 검증 결과

## 1. 요구사항 매핑 결과
{mapping_result}

## 2. 누락 항목 분석
{missing_items}

        """
        
        # 다운로드 데이터 미리 생성
        quality_download_data = create_quality_download_data("rfp_quality_check.docx", combined_content)
        
        # HTML 다운로드 링크 사용 (페이지 리로드 방지)
        quality_link = create_quality_download_link(quality_download_data, "rfp_quality_check.docx", "📋 품질 검증 결과 다운로드")
        st.markdown(quality_link, unsafe_allow_html=True)
                
    except Exception as e:
        st.error(f"품질 검증 결과 생성 중 오류: {str(e)}")


def generate_requirements_mapping_with_urls(azure_services, rfp_info, main_rfp_url, main_proposal_url):
    """URL 기반 요구사항 매핑 생성"""
    
    messages = [
        {
            "role": "system",
            "content": f"""당신은 {rfp_info.get('industry', '금융')} 업종의 제안서 품질 관리 전문가입니다.
            첨부된 RFP와 제안서를 분석하여 요구사항 매핑을 수행해주세요."""
        },
        {
            "role": "user",
            "content": f"""
            업종: {rfp_info.get('industry', '금융')}
            RFP 정보: {rfp_info.get('name', '')}
            
            첨부 파일:
            1. RFP 파일: {main_rfp_url}
            2. 제안서 파일: {main_proposal_url}
            
            다음 형식으로 요구사항 매핑을 분석해주세요:
            
            ## 요구사항 매핑 분석
            
            ### 1. 완전히 충족된 요구사항
            - [요구사항]: [제안서에서의 응답]
            - 매핑도: 100%
            
            ### 2. 부분적으로 충족된 요구사항
            - [요구사항]: [제안서에서의 응답]
            - 매핑도: [%]
            - 개선 필요사항: [구체적 개선 방안]
            
            ### 3. 충족되지 않은 요구사항
            - [요구사항]: [누락된 내용]
            - 매핑도: 0%
            - 보완 방안: [구체적 보완 방안]
            
            ### 4. 전체 매핑 점수
            - 전체 충족도: [%]
            - 우선 개선 항목: [상위 3개]
            """
        }
    ]
    
    return azure_services.call_openai_with_files(messages, [main_rfp_url, main_proposal_url])



def detect_missing_items_with_urls(azure_services, rfp_info, main_rfp_url, main_proposal_url):
    """URL 기반 누락 항목 자동 감지"""
    
    messages = [
        {
            "role": "system",
            "content": f"""당신은 {rfp_info.get('industry', '금융')} 업종의 제안서 검토 전문가입니다.
            첨부된 RFP와 제안서를 분석하여 누락된 항목을 체계적으로 감지해주세요."""
        },
        {
            "role": "user",
            "content": f"""
            업종: {rfp_info.get('industry', '금융')}
            
            첨부 파일:
            1. RFP 파일: {main_rfp_url}
            2. 제안서 파일: {main_proposal_url}
            
            다음 관점에서 누락된 항목을 분석해주세요:
            
            ## 누락 항목 분석
            
            ### 1. 기능적 요구사항 누락
            - [누락된 기능]: [상세 설명]
              - 중요도: [High/Medium/Low] (점수: [N/10점])
              - 영향 범위: [전체/부분적/제한적]
              - 누락 리스크: [N/10점]
              - 보완 방안: [구체적 제안]
            
            ### 2. 비기능적 요구사항 누락
            - [누락된 요구사항]: [상세 설명]
              - 중요도: [High/Medium/Low] (점수: [N/10점])
              - 시스템 영향도: [N/10점]
              - 누락 리스크: [N/10점]
              - 보완 방안: [구체적 제안]
            
            ### 3. 기술적 요구사항 누락
            - [누락된 기술]: [상세 설명]
              - 중요도: [High/Medium/Low] (점수: [N/10점])
              - 기술적 리스크: [N/10점]
              - 보완 난이도: [상/중/하]
              - 보완 방안: [구체적 제안]
            
            ### 4. 비즈니스 요구사항 누락
            - [누락된 요구사항]: [상세 설명]
              - 중요도: [High/Medium/Low] (점수: [N/10점])
              - 비즈니스 영향: [N/10점]
              - 긴급도: [즉시/단기/중장기]
              - 보완 방안: [구체적 제안]
            
            ### 5. 우선순위별 개선 계획
            - **High 우선순위 (8-10점)**: 즉시 보완 필요, 프로젝트 성공에 필수
            - **Medium 우선순위 (5-7점)**: 단계적 보완, 품질 향상에 중요
            - **Low 우선순위 (1-4점)**: 선택적 보완, 추가 가치 제공
            """
        }
    ]
    
    return azure_services.call_openai_with_files(messages, [main_rfp_url, main_proposal_url])

def generate_requirements_mapping_manual(azure_services, rfp_info, proposal_content):
    """수동 입력 기반 요구사항 매핑 생성"""
    
    messages = [
        {
            "role": "system",
            "content": f"""당신은 {rfp_info.get('industry', '금융')} 업종의 제안서 품질 관리 전문가입니다.
            RFP와 제안서를 분석하여 요구사항 매핑을 수행해주세요."""
        },
        {
            "role": "user",
            "content": f"""
            업종: {rfp_info.get('industry', '금융')}
            RFP 정보: {rfp_info.get('name', '')}
            RFP 요약: {rfp_info.get('rfp_summary', '')}
            RFP 내용: {rfp_info.get('rfp_content', '')}
            
            제안서 내용: {proposal_content}
            
            다음 형식으로 요구사항 매핑을 분석하고, **구체적인 수치와 평가**를 제공해주세요:
            
            ## 요구사항 매핑 분석
            
            ### 1. 완전히 충족된 요구사항
            - [요구사항]: [제안서에서의 응답]
              - 매핑도: 100% (충족도 점수: 10/10점)
              - 제안 품질: [상/중/하]
              - 차별화 포인트: [있음/보통/없음]
            
            ### 2. 부분적으로 충족된 요구사항
            - [요구사항]: [제안서에서의 응답]
              - 매핑도: [N%] (충족도 점수: [N/10점])
              - 충족된 부분: [구체적 내용]
              - 부족한 부분: [구체적 내용]
              - 개선 필요도: [상/중/하]
              - 개선 필요사항: [구체적 개선 방안]
            
            ### 3. 충족되지 않은 요구사항
            - [요구사항]: [누락된 내용]
              - 매핑도: 0% (충족도 점수: 0/10점)
              - 중요도: [N/10점]
              - 누락 리스크: [상/중/하]
              - 보완 방안: [구체적 보완 방안]
              - 예상 보완 시간: [N 일]
            
            ### 4. 전체 매핑 점수
            - 전체 충족도: [N%] (평균 점수: [N/10점])
            - 완전 충족: [N개 항목] ([N%])
            - 부분 충족: [N개 항목] ([N%])
            - 미충족: [N개 항목] ([N%])
            - 우선 개선 항목: [상위 3개, 중요도 순]
            
            **매핑도 평가 기준:**
            - 90-100%: 완전 충족, 우수한 제안
            - 70-89%: 대부분 충족, 일부 보완 필요
            - 50-69%: 부분 충족, 상당한 개선 필요
            - 30-49%: 불충분, 대폭 보완 필요
            - 0-29%: 거의 미충족, 전면 재작성 필요
            """
        }
    ]
    
    return azure_services.call_openai(messages)

def detect_missing_items_manual(azure_services, rfp_info, proposal_content):
    """수동 입력 기반 누락 항목 자동 감지"""
    
    messages = [
        {
            "role": "system",
            "content": f"""당신은 {rfp_info.get('industry', '금융')} 업종의 제안서 검토 전문가입니다.
            RFP와 제안서를 분석하여 누락된 항목을 체계적으로 감지해주세요."""
        },
        {
            "role": "user",
            "content": f"""
            업종: {rfp_info.get('industry', '금융')}
            RFP 요약: {rfp_info.get('rfp_summary', '')}
            RFP 내용: {rfp_info.get('rfp_content', '')}
            
            제안서 내용: {proposal_content}
            
            다음 관점에서 누락된 항목을 분석해주세요:
            
            ## 누락 항목 분석
            
            ### 1. 기능적 요구사항 누락
            - [누락된 기능]: [상세 설명]
              - 중요도: [High/Medium/Low] (점수: [N/10점])
              - 영향 범위: [전체/부분적/제한적]
              - 누락 리스크: [N/10점]
              - 보완 방안: [구체적 제안]
            
            ### 2. 비기능적 요구사항 누락
            - [누락된 요구사항]: [상세 설명]
              - 중요도: [High/Medium/Low] (점수: [N/10점])
              - 시스템 영향도: [N/10점]
              - 누락 리스크: [N/10점]
              - 보완 방안: [구체적 제안]
            
            ### 3. 기술적 요구사항 누락
            - [누락된 기술]: [상세 설명]
              - 중요도: [High/Medium/Low] (점수: [N/10점])
              - 기술적 리스크: [N/10점]
              - 보완 난이도: [상/중/하]
              - 보완 방안: [구체적 제안]
            
            ### 4. 비즈니스 요구사항 누락
            - [누락된 요구사항]: [상세 설명]
              - 중요도: [High/Medium/Low] (점수: [N/10점])
              - 비즈니스 영향: [N/10점]
              - 긴급도: [즉시/단기/중장기]
              - 보완 방안: [구체적 제안]
            
            ### 5. 우선순위별 개선 계획
            - **High 우선순위 (8-10점)**: 즉시 보완 필요, 프로젝트 성공에 필수
            - **Medium 우선순위 (5-7점)**: 단계적 보완, 품질 향상에 중요
            - **Low 우선순위 (1-4점)**: 선택적 보완, 추가 가치 제공
            """
        }
    ]
    
    return azure_services.call_openai(messages)

def save_quality_results_to_directory(mapping_result, missing_items):
    """품질 검증 결과를 별도 디렉토리에 자동 저장"""
    try:
        azure_services = st.session_state.azure_services
        
        # 세션 상태에서 현재 작업 디렉토리 가져오기
        if not hasattr(st.session_state, 'current_directory') or not st.session_state.current_directory:
            # 디렉토리 정보가 없으면 저장을 건너뛰고 조용히 종료
            return
        
        # 기존 RFP 분석 디렉토리 사용 (같은 디렉토리 유지)
        container_name = st.session_state.current_container
        directory_name = st.session_state.current_directory
        
        # 타임스탬프 생성
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # 통합 품질 검증 결과 생성
        combined_content = f"""
# 제안서 품질 검증 결과

## 1. 요구사항 매핑 결과
{mapping_result}

## 2. 누락 항목 분석
{missing_items}

        """
        
        # Word 문서 생성 및 저장
        doc = Document()
        doc.add_heading('제안서 품질 검증 결과', 0)
        doc.add_paragraph(combined_content)
        
        # 임시 파일로 저장
        temp_filename = "temp_rfp_quality_check.docx"
        doc.save(temp_filename)
        
        with open(temp_filename, 'rb') as f:
            file_data = f.read()
        
        # Azure에 업로드 (타임스탬프 추가하여 고유 파일명 생성)
        unique_filename = f"rfp_quality_check_{timestamp}.docx"
        azure_services.upload_file_to_directory(container_name, directory_name, unique_filename, file_data)
        
        # 임시 파일 삭제
        import os
        os.remove(temp_filename)
        
        st.success(f"📁 품질 검증 결과가 디렉토리에 저장되었습니다: {unique_filename}")
        
    except Exception as e:
        st.error(f"품질 검증 결과 저장 중 오류: {str(e)}")

def create_quality_download_data(filename, content):
    """품질 검증 결과 다운로드용 데이터 생성"""
    try:
        # Word 문서 생성
        doc = Document()
        doc.add_heading('제안서 품질 검증 결과', 0)
        doc.add_paragraph(content)
        
        # 임시 파일로 저장
        temp_filename = f"temp_{filename}"
        doc.save(temp_filename)
        
        # 파일 읽기
        with open(temp_filename, 'rb') as f:
            file_data = f.read()
        
        # 임시 파일 삭제
        import os
        os.remove(temp_filename)
        
        return file_data
        
    except Exception as e:
        st.error(f"다운로드 데이터 생성 중 오류: {str(e)}")
        return b""

def create_quality_download_link(data, filename, label):
    """품질 검증 다운로드 링크 생성 (페이지 리로드 방지)"""
    import base64
    
    # 데이터를 base64로 인코딩
    b64 = base64.b64encode(data).decode()
    
    # HTML 다운로드 링크 생성
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" style="display: inline-block; padding: 0.5rem 1rem; background-color: #1f77b4; color: white; text-decoration: none; border-radius: 0.25rem; border: none; cursor: pointer;">{label}</a>'
    
    return href

def extract_text_from_pdf_bytes(file_bytes):
    """PDF 바이트 데이터에서 텍스트 추출"""
    try:
        # pdfplumber를 사용한 텍스트 추출 (더 정확함)
        text_content = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n"
        
        # pdfplumber로 텍스트를 추출하지 못한 경우 PyPDF2 사용
        if not text_content.strip():
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
        
        return text_content.strip()
        
    except Exception as e:
        st.error(f"PDF 바이트 데이터 텍스트 추출 중 오류: {str(e)}")
        return ""

def extract_text_from_docx_bytes(file_bytes):
    """DOCX 바이트 데이터에서 텍스트 추출"""
    try:
        # BytesIO를 사용하여 Document 객체 생성
        doc = Document(io.BytesIO(file_bytes))
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        return text_content.strip()
        
    except Exception as e:
        st.error(f"DOCX 바이트 데이터 텍스트 추출 중 오류: {str(e)}")
        return ""


