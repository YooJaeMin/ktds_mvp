"""
RFP 분석 페이지
"""
import streamlit as st
import json
import time
from datetime import datetime
from docx import Document
import pandas as pd
import PyPDF2
import pdfplumber
import io

def show():
    """RFP 분석 페이지 표시"""
    st.title("🔍 RFP 분석")
    st.markdown("RFP 문서를 분석하여 요구사항을 추출하고 분류합니다.")
    
    # 분석 모드 선택
    analysis_mode = st.radio(
        "분석 모드 선택",
        ["새 RFP 문서 분석", "저장된 RFP 재분석"],
        horizontal=True
    )
    
    if analysis_mode == "새 RFP 문서 분석":
        show_new_rfp_analysis()
    else:
        show_stored_rfp_analysis()

def show_new_rfp_analysis():
    """새 RFP 문서 분석"""
    st.subheader("📄 새 RFP 문서 분석")
    
    # 파일 업로드
    uploaded_file = st.file_uploader(
        "RFP 문서 업로드",
        type=['pdf', 'docx', 'txt'],
        help="분석할 RFP 문서를 업로드하세요. PDF, DOCX, TXT 형식을 지원합니다."
    )
    
    if uploaded_file:
        # 파일 정보 표시
        st.info(f"📁 업로드된 파일: {uploaded_file.name} ({uploaded_file.size} bytes)")
        
        # 분석 옵션
        st.subheader("🔧 분석 옵션")
        
        col1, col2 = st.columns(2)
        
        with col1:
            industry = st.selectbox(
                "업종",
                ["은행", "보험", "증권", "카드", "기타"],
                help="RFP의 대상 업종을 선택하세요."
            )
            
            analysis_depth = st.selectbox(
                "분석 깊이",
                ["기본", "상세", "심화"],
                help="분석의 상세도를 선택하세요."
            )
        
        with col2:
            focus_area = st.multiselect(
                "중점 분석 영역",
                ["기능 요구사항", "비기능 요구사항", "기술 스택", "보안 요구사항", "성능 요구사항", "통합 요구사항"],
                default=["기능 요구사항", "비기능 요구사항"],
                help="중점적으로 분석할 영역을 선택하세요."
            )
        
        # 분석 실행
        if st.button("분석 시작", type="primary"):
            if not focus_area:
                st.error("최소 하나의 중점 분석 영역을 선택해주세요.")
                return
            
            with st.spinner("RFP 문서를 분석하고 있습니다..."):
                analyze_rfp_document(uploaded_file, industry, analysis_depth, focus_area)
            
            st.success("RFP 분석이 완료되었습니다!")

def show_stored_rfp_analysis():
    """저장된 RFP 재분석"""
    st.subheader("📁 저장된 RFP 재분석")
    
    try:
        azure_services = st.session_state.azure_services
        directories = azure_services.get_directories()
        
        if not directories:
            st.warning("저장된 RFP가 없습니다. 먼저 RFP를 업로드하고 분석해주세요.")
            return
        
        # 디렉토리 목록을 테이블로 표시
        st.subheader("📋 저장된 RFP 목록")
        
        # 페이지네이션 설정
        items_per_page = 10
        total_pages = (len(directories) + items_per_page - 1) // items_per_page
        
        if total_pages > 1:
            page = st.selectbox("페이지 선택", range(1, total_pages + 1), key="rfp_page")
        else:
            page = 1
        
        start_idx = (page - 1) * items_per_page
        end_idx = start_idx + items_per_page
        current_directories = directories[start_idx:end_idx]
        
        # 테이블 데이터 생성 (디렉토리명 및 선택 컬럼 제외)
        table_data = []
        for i, directory in enumerate(current_directories):
            table_data.append({
                "한글명": directory['korean_name'],
                "생성일": directory['created_date'],
                "프로젝트 요약": directory['project_summary'][:50] + "..." if len(directory['project_summary']) > 50 else directory['project_summary']
            })
        
        # 테이블 표시
        df = pd.DataFrame(table_data)
        st.dataframe(df, width='stretch')
        
        # 라디오 버튼으로 선택 (테이블과 함께 표시)
        col1, col2 = st.columns([1, 3])
        with col1:
            st.markdown("**재분석할 RFP 선택:**")
        with col2:
            selected_index = st.radio(
                "RFP 선택",
                range(len(current_directories)),
                format_func=lambda x: f"{current_directories[x]['korean_name']}",
                horizontal=True,
                key="rfp_reanalysis_selection",
                label_visibility="collapsed"
            )
        
        if selected_index is not None:
            selected_directory = current_directories[selected_index]
            
            # 파일 목록 표시
            container_name = "rfp-documents"
            files = azure_services.list_files_in_directory(container_name, selected_directory['name'])
            if files:
                st.info(f"📄 {selected_directory['korean_name']}에 저장된 파일: {', '.join(files)}")
                
                # 재분석 옵션
                st.subheader("🔄 재분석 옵션")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    new_industry = st.selectbox(
                        "업종 (변경 가능)",
                        ["은행", "보험", "증권", "카드", "기타"],
                        help="업종을 변경하여 재분석할 수 있습니다."
                    )
                    
                    new_depth = st.selectbox(
                        "분석 깊이 (변경 가능)",
                        ["기본", "상세", "심화"],
                        help="분석 깊이를 변경하여 재분석할 수 있습니다."
                    )
                
                with col2:
                    new_focus = st.multiselect(
                        "중점 분석 영역 (변경 가능)",
                        ["기능 요구사항", "비기능 요구사항", "기술 스택", "보안 요구사항", "성능 요구사항", "통합 요구사항"],
                        default=["기능 요구사항", "비기능 요구사항"],
                        help="중점 분석 영역을 변경할 수 있습니다."
                    )
                
                if st.button("재분석 시작", type="primary"):
                    if not new_focus:
                        st.error("최소 하나의 중점 분석 영역을 선택해주세요.")
                        return
                    
                    with st.spinner("저장된 RFP를 재분석하고 있습니다..."):
                        reanalyze_stored_rfp(selected_directory['name'], new_industry, new_depth, new_focus)
                    
                    st.success("RFP 재분석이 완료되었습니다!")
            else:
                st.warning("선택된 디렉토리에 파일이 없습니다.")
                
    except Exception as e:
        st.error(f"오류가 발생했습니다: {str(e)}")

def analyze_rfp_document(uploaded_file, industry, analysis_depth, focus_area):
    """RFP 문서 분석 실행"""
    try:
        # 새로운 분석 시작 시 세션 상태 초기화
        if hasattr(st.session_state, 'current_directory'):
            del st.session_state.current_directory
        if hasattr(st.session_state, 'current_container'):
            del st.session_state.current_container
        
        # 파일에서 텍스트 추출
        st.info("📄 파일에서 텍스트를 추출하고 있습니다...")
        content = extract_text_from_uploaded_file(uploaded_file)
        
        if not content:
            st.error("파일에서 텍스트를 추출할 수 없습니다. 파일이 손상되었거나 지원하지 않는 형식일 수 있습니다.")
            return
        
        # 추출된 텍스트 미리보기
        st.subheader("📋 추출된 텍스트 미리보기")
        preview_length = 500
        if len(content) > preview_length:
            st.text_area("텍스트 미리보기 (처음 500자)", content[:preview_length] + "...", height=200, disabled=True)
            st.info(f"전체 텍스트 길이: {len(content)}자")
        else:
            st.text_area("추출된 텍스트", content, height=200, disabled=True)
        
        # 1단계: RFP 파일을 먼저 Azure Storage에 업로드
        st.info("📤 RFP 파일을 Azure Storage에 업로드하고 있습니다...")
        uploaded_file.seek(0)  # 파일 포인터를 처음으로 리셋
        file_content = uploaded_file.read()
        
        # 파일 크기 정보 표시 (디버깅용)
        st.info(f"📊 원본 파일 크기: {len(file_content)} bytes, 추출된 텍스트 길이: {len(content)}자")
        
        # RFP 파일 업로드 (디렉토리 생성 및 세션 상태 설정)
        save_to_azure_storage(uploaded_file.name, file_content, analysis_depth, focus_area, extracted_text=content)
        
        # 2단계: 분석 결과 생성 (같은 디렉토리에 저장)
        st.info("🔍 텍스트를 분석하고 있습니다...")
        generate_analysis_results(content, industry, analysis_depth, focus_area, uploaded_file.name)
        
    except Exception as e:
        st.error(f"분석 중 오류가 발생했습니다: {str(e)}")

def reanalyze_stored_rfp(directory_name, industry, analysis_depth, focus_area):
    """저장된 RFP 재분석"""
    try:
        # 재분석 시작 시 세션 상태 초기화
        if hasattr(st.session_state, 'current_directory'):
            del st.session_state.current_directory
        if hasattr(st.session_state, 'current_container'):
            del st.session_state.current_container
        
        azure_services = st.session_state.azure_services
        container_name = "rfp-documents"
        files = azure_services.list_files_in_directory(container_name, directory_name)
        
        # 메타데이터에서 한글명 가져오기
        metadata = azure_services.get_directory_metadata_from_path(container_name, directory_name)
        korean_name = metadata.get('korean_name', directory_name)
        
        ##시스템 출력
        st.info(f"재분석할 RFP: {korean_name}")
        if not files:
            st.error("재분석할 파일이 없습니다.")
            return
        
        # main_rfp_로 시작하는 파일 찾기
        main_rfp_file = None
        for file_name in files:
            if file_name.startswith('main_rfp_'):
                main_rfp_file = file_name
                break
        
        if not main_rfp_file:
            st.error("main_rfp_ 파일을 찾을 수 없습니다.")
            return
        
        container_name = "rfp-documents"
        file_content = azure_services.download_file_from_directory(container_name, directory_name, main_rfp_file)
        
        if file_content:
            # 파일 확장자에 따라 적절한 텍스트 추출 방법 선택
            if main_rfp_file.lower().endswith('.pdf'):
                # PDF 파일의 경우 텍스트 추출 함수 사용
                content = extract_text_from_pdf_bytes(file_content)
            elif main_rfp_file.lower().endswith('.docx'):
                # DOCX 파일의 경우 텍스트 추출 함수 사용
                content = extract_text_from_docx_bytes(file_content)
            elif main_rfp_file.lower().endswith('.txt'):
                # TXT 파일의 경우 UTF-8 디코딩
                content = file_content.decode('utf-8', errors='ignore')
            else:
                # 기본적으로 UTF-8 디코딩 시도
                content = file_content.decode('utf-8', errors='ignore')
            
            # 새로운 디렉토리 생성 (재분석 결과용) - 영어와 숫자만 사용
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            new_directory_name = f"rfpreanalysis{timestamp}"
            
            # rfp-documents 컨테이너 내에 새 디렉토리 생성
            container_name = "rfp-documents"
            
            # 원본 RFP 파일을 새 디렉토리에 복사
            if azure_services.upload_file_to_directory(container_name, new_directory_name, main_rfp_file, file_content):
                # 현재 작업 디렉토리를 세션 상태에 저장
                st.session_state.current_directory = new_directory_name
                st.session_state.current_container = container_name
                
                # 분석 결과 생성 및 표시 (자동 저장 비활성화)
                generate_analysis_results(content, industry, analysis_depth, focus_area, main_rfp_file, auto_save=False)
                
                # 재분석 결과를 새 디렉토리에 저장
                save_reanalysis_results(container_name, new_directory_name, content, industry, analysis_depth, focus_area)
                
                st.success(f"재분석 결과가 새 디렉토리에 저장되었습니다: {new_directory_name}")
            else:
                st.error("새 디렉토리 생성에 실패했습니다.")
        else:
            st.error("파일을 다운로드할 수 없습니다.")
            
    except Exception as e:
        st.error(f"재분석 중 오류가 발생했습니다: {str(e)}")

def save_reanalysis_results(container_name, directory_name, content, industry, analysis_depth, focus_area):
    """재분석 결과를 새 디렉토리에 저장"""
    try:
        azure_services = st.session_state.azure_services
        
        # 실제 RFP 내용을 분석하여 프로젝트 요약 생성
        project_summary = generate_enhanced_project_summary(content, industry, analysis_depth, focus_area)
        korean_name = generate_korean_project_name(project_summary)
        
        # 재분석 메타데이터 생성
        metadata = {
            'korean_name': f"{korean_name} (재분석)",
            'created_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'project_summary': project_summary,
            'original_filename': "reanalysis_result",
            'analysis_depth': analysis_depth,
            'focus_areas': focus_area,
            'is_reanalysis': True
        }
        
        azure_services.save_directory_metadata_to_path(container_name, directory_name, metadata)
        
        # 재분석 결과 DOCX 파일들 생성 및 저장
        save_reanalysis_docx_files(container_name, directory_name, content, industry, analysis_depth, focus_area)
        
    except Exception as e:
        st.error(f"재분석 결과 저장 중 오류: {str(e)}")

def save_reanalysis_docx_files(container_name, directory_name, content, industry, analysis_depth, focus_area):
    """재분석 결과 DOCX 파일들을 생성하고 저장"""
    try:
        azure_services = st.session_state.azure_services
        
        # 병렬 처리로 분석 결과 생성
        import concurrent.futures
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            # 모든 분석을 병렬로 실행
            future_requirements = executor.submit(extract_requirements_with_azure, azure_services, content, industry, analysis_depth, focus_area)
            future_keywords = executor.submit(analyze_keywords_with_azure, azure_services, content, industry, analysis_depth)
            future_summary = executor.submit(generate_summary_report_with_azure, azure_services, content, industry, analysis_depth, focus_area)
            
            # 결과 수집
            requirements = future_requirements.result()
            keywords = future_keywords.result()
            summary = future_summary.result()
        
        # 통합 분석 결과 생성
        combined_content = f"""
# RFP 재분석 결과

## 1. 요구사항 추출
{requirements}

## 2. 키워드 분석
{keywords}

## 3. 요약 보고서
{summary}
        """
        
        # 1. 상세 분석 결과 저장
        doc = Document()
        doc.add_heading('RFP 재분석 결과', 0)
        doc.add_paragraph(combined_content)
        
        # 임시 파일로 저장
        temp_detailed = "temp_reanalysis_result_detail.docx"
        doc.save(temp_detailed)
        
        with open(temp_detailed, 'rb') as f:
            detailed_data = f.read()
        
        # 타임스탬프 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Azure에 업로드 (타임스탬프 포함)
        detailed_filename = f"analysis_result_detail_{timestamp}.docx"
        azure_services.upload_file_to_directory(container_name, directory_name, detailed_filename, detailed_data)
        
        # 2. 요약 보고서 저장
        doc_summary = Document()
        doc_summary.add_heading('RFP 재분석 요약 보고서', 0)
        doc_summary.add_paragraph(summary)
        
        # 임시 파일로 저장
        temp_summary = "temp_reanalysis_result_summary.docx"
        doc_summary.save(temp_summary)
        
        with open(temp_summary, 'rb') as f:
            summary_data = f.read()
        
        # Azure에 업로드 (타임스탬프 포함)
        summary_filename = f"analysis_result_summary_{timestamp}.docx"
        azure_services.upload_file_to_directory(container_name, directory_name, summary_filename, summary_data)
        
        # 임시 파일 삭제
        import os
        os.remove(temp_detailed)
        os.remove(temp_summary)
        
        st.success("📁 재분석 결과 DOCX 파일들이 디렉토리에 저장되었습니다!")
        
    except Exception as e:
        st.error(f"재분석 DOCX 파일 저장 중 오류: {str(e)}")

def generate_analysis_results(content, industry, analysis_depth, focus_area, file_name, auto_save=True):
    """분석 결과 생성 및 표시"""
    try:
        azure_services = st.session_state.azure_services
        
        
        # 분석 결과 다운로드 안내 (탭 위쪽에)
        st.markdown("---")
        with st.container():
            st.markdown("### 📥 분석 결과 다운로드")
            st.info("💡 **분석이 완료되면 아래 탭에서 결과를 확인하고, 페이지 하단의 다운로드 버튼을 통해 결과를 저장할 수 있습니다.**")
        
        # 탭으로 결과 표시
        tab1, tab2, tab3 = st.tabs([
            "📋 요구사항 추출", 
            "🔍 키워드 분석",
            "📄 요약 보고서"
        ])
        
        # 병렬 처리를 위한 placeholder 생성
        with tab1:
            req_placeholder = st.empty()
            req_placeholder.info("🔄 요구사항 추출 중...")
        
        with tab2:
            keyword_placeholder = st.empty()
            keyword_placeholder.info("🔄 키워드 분석 중...")
        
        with tab3:
            summary_placeholder = st.empty()
            summary_placeholder.info("🔄 요약 보고서 생성 중...")
        
        # Azure 서비스를 미리 가져와서 병렬 처리에서 사용할 수 있도록 준비
        azure_services = st.session_state.azure_services
        
        # 최적화된 병렬 처리 실행
        from modules.performance import parallel_analysis_executor
        
        analyses = [
            {
                'name': 'requirements',
                'func': extract_requirements_with_azure,
                'args': (azure_services, content, industry, analysis_depth, focus_area),
                'kwargs': {}
            },
            {
                'name': 'keywords', 
                'func': analyze_keywords_with_azure,
                'args': (azure_services, content, industry, analysis_depth),
                'kwargs': {}
            },
            {
                'name': 'summary',
                'func': generate_summary_report_with_azure,
                'args': (azure_services, content, industry, analysis_depth, focus_area),
                'kwargs': {}
            }
        ]
        
        def run_analysis_with_azure(azure_services, content, industry, focus_area, analysis_depth):
            """최적화된 Azure 서비스 분석 실행"""
            results = parallel_analysis_executor(analyses, max_workers=3)
            
            # 결과를 순차적으로 표시 (완료되는 대로)
            # 요구사항 추출 결과 표시
            requirements = results['requirements']
            req_placeholder.markdown(
                f"""
                <div style="max-height: 600px; overflow-y: auto; padding: 15px; border: 1px solid #e0e0e0; border-radius: 8px; background-color: #fafafa;">
                    {requirements}
                </div>
                """, 
                unsafe_allow_html=True
            )
            
            # 키워드 분석 결과 표시
            keywords = results['keywords']
            keyword_placeholder.markdown(
                f"""
                <div style="max-height: 600px; overflow-y: auto; padding: 15px; border: 1px solid #e0e0e0; border-radius: 8px; background-color: #fafafa;">
                    {keywords}
                </div>
                """, 
                unsafe_allow_html=True
            )
            create_keyword_cloud()
            
            # 요약 보고서 결과 표시
            summary = results['summary']
            summary_placeholder.markdown(
                f"""
                <div style="max-height: 600px; overflow-y: auto; padding: 15px; border: 1px solid #e0e0e0; border-radius: 8px; background-color: #fafafa;">
                    {summary}
                </div>
                """, 
                unsafe_allow_html=True
            )
            
            return results
        
        # 분석 실행
        results = run_analysis_with_azure(azure_services, content, industry, focus_area, analysis_depth)
        requirements = results['requirements']
        keywords = results['keywords']
        summary = results['summary']
        
        # 분석 결과를 디렉토리에 자동 저장 (auto_save가 True일 때만)
        if auto_save:
            save_analysis_results_to_directory(content, industry, analysis_depth, focus_area, requirements, keywords, summary)
        
        # 모든 분석 결과를 하나로 통합
        combined_content = f"""
# RFP 상세 분석 결과

## 1. 요구사항 추출 결과
{requirements}

## 2. 키워드 분석 결과
{keywords}
        """
        
        # 파일명 생성
        detailed_filename = "analysis_result_detail.docx"
        summary_filename = "analysis_result_summary.docx"
        
        # 세션 상태에 결과 저장
        st.session_state.requirements = requirements
        st.session_state.keywords = keywords
        st.session_state.summary = summary
        st.session_state.detailed_filename = detailed_filename
        st.session_state.summary_filename = summary_filename
        st.session_state.combined_content = combined_content
        
        # 다운로드 버튼들을 탭 아래쪽에 표시
        st.markdown("---")
        
        # 성공 메시지와 다운로드 섹션을 강조
        with st.container():
            st.success("🎉 **분석이 완료되었습니다!** 아래 버튼을 통해 결과를 다운로드하세요.")
            st.markdown("### 📥 분석 결과 다운로드")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # 다운로드 데이터 미리 생성
            detailed_filename = st.session_state.get('detailed_filename', 'analysis_result_detail.docx')
            detailed_download_data = create_download_data(detailed_filename, combined_content)
            
            # HTML 다운로드 링크 사용 (페이지 리로드 방지)
            detailed_link = create_download_link(detailed_download_data, detailed_filename, "📄 상세 분석 결과 다운로드 (DOCX)")
            st.markdown(detailed_link, unsafe_allow_html=True)
        
        with col2:
            # 다운로드 데이터 미리 생성
            summary_filename = st.session_state.get('summary_filename', 'analysis_result_summary.docx')
            summary_download_data = create_download_data(summary_filename, summary)
            
            # HTML 다운로드 링크 사용 (페이지 리로드 방지)
            summary_link = create_download_link(summary_download_data, summary_filename, "📋 요약 보고서 다운로드 (DOCX)")
            st.markdown(summary_link, unsafe_allow_html=True)
                
    except Exception as e:
        st.error(f"분석 결과 생성 중 오류: {str(e)}")

def extract_requirements(content, industry, analysis_depth, focus_area):
    """요구사항 추출"""
    azure_services = st.session_state.azure_services
    return extract_requirements_with_azure(azure_services, content, industry, analysis_depth, focus_area)

def extract_requirements_with_azure(azure_services, content, industry, analysis_depth, focus_area):
    """Azure 서비스를 전달받아 요구사항 추출"""
    
    # 분석 깊이에 따른 지시사항
    depth_instructions = {
        "기본": "주요 요구사항 총 3개만 간략하게 추출하세요. 각 항목당 1-2줄로 요약합니다.",
        "상세": "모든 요구사항별로 5개를 상세하게 추출하세요. 각 요구사항의 배경과 목적을 포함합니다.",
        "심화": "모든 요구사항별로 10개를 매우 상세하게 추출하고, 잠재적 이슈, 구현 고려사항, 관련 요구사항과의 연관성까지 분석하세요."
    }
    
    messages = [
        {
            "role": "system",
            "content": f"당신은 {industry} 업종의 RFP 분석 전문가입니다. RFP 문서에서 요구사항을 체계적으로 추출해주세요. 분석 깊이: {analysis_depth}"
        },
        {
            "role": "user",
            "content": f"""
            업종: {industry}
            분석 깊이: {analysis_depth}
            중점 분석 영역: {', '.join(focus_area)}
            
            ** 분석 지침: {depth_instructions.get(analysis_depth, depth_instructions['기본'])} **
            
            RFP 문서 내용:
            {content}
            
            다음 형식으로 요구사항을 추출하고, **구체적인 평가 점수와 기준**을 명시해주세요:
            
            ## 요구사항 추출 결과
            
            ### 1. 기능적 요구사항
            - [요구사항 1]: [상세 설명]
              - 우선순위: [High/Medium/Low] (점수: [N/10점])
              - 구현 난이도: [상/중/하] (예상 공수: [N 인일])
              - 비즈니스 가치: [N/10점]
            {"  - 구현 고려사항: [상세 설명]" if analysis_depth == "심화" else ""}
            
            ### 2. 비기능적 요구사항
            - [요구사항 1]: [상세 설명]
              - 중요도: [Critical/Important/Nice-to-have] (점수: [N/10점])
              - 기술적 난이도: [상/중/하]
              - 성능 목표: [구체적 수치, 예: 응답시간 < 2초]
            {"  - 기술적 제약사항: [상세 설명]" if analysis_depth == "심화" else ""}
            
            ### 3. 기술적 요구사항
            - [요구사항 1]: [상세 설명]
              - 복잡도: [High/Medium/Low] (점수: [N/10점])
              - 기술 성숙도: [검증됨/보통/신기술]
              - 예상 리스크: [N/10점]
            {"  - 기술 스택 권장사항: [상세 설명]" if analysis_depth in ["상세", "심화"] else ""}
            
            ### 4. 비즈니스 요구사항
            - [요구사항 1]: [상세 설명]
              - 비즈니스 임팩트: [High/Medium/Low] (점수: [N/10점])
              - 긴급도: [즉시/단기/중장기]
              - 투자 대비 효과: [N/10점]
            {"  - ROI 분석: [상세 설명]" if analysis_depth == "심화" else ""}
            
            **평가 기준:**
            - 10점: 프로젝트 필수 요소, 최우선 처리
            - 8-9점: 핵심 요구사항, 높은 우선순위
            - 6-7점: 중요 요구사항, 중간 우선순위
            - 4-5점: 일반 요구사항, 필요시 조정 가능
            - 1-3점: 선택적 요구사항, 추가 기능
            """
        }
    ]
    
    return azure_services.call_openai(messages)

def analyze_keywords(content, industry, analysis_depth):
    """키워드 분석"""
    azure_services = st.session_state.azure_services
    return analyze_keywords_with_azure(azure_services, content, industry, analysis_depth)

def analyze_keywords_with_azure(azure_services, content, industry, analysis_depth):
    """Azure 서비스를 전달받아 키워드 분석"""
    
    # 분석 깊이에 따른 키워드 분석 범위
    depth_instructions = {
        "기본": "상위 5개 주요 키워드만 각 카테고리별로 추출하세요.",
        "상세": "상위 10개 키워드를 추출하고, 각 키워드의 문맥과 의미를 분석하세요.",
        "심화": "상위 15개 키워드를 추출하고, 키워드 간 연관성, 트렌드 분석, 산업별 특이사항까지 상세히 분석하세요."
    }
    
    messages = [
        {
            "role": "system",
            "content": f"당신은 {industry} 업종의 텍스트 분석 전문가입니다. RFP 문서의 키워드를 분석해주세요. 분석 깊이: {analysis_depth}"
        },
        {
            "role": "user",
            "content": f"""
            업종: {industry}
            분석 깊이: {analysis_depth}
            
            ** 분석 지침: {depth_instructions.get(analysis_depth, depth_instructions['기본'])} **
            
            RFP 문서 내용:
            {content}
            
            다음 관점에서 키워드를 분석하고, 반드시 **구체적인 수치**를 포함해주세요:
            
            ## 키워드 분석 결과
            
            ### 1. 기술 키워드
            - [키워드 1]: 출현 빈도 [N회], 중요도 [N/10점], 문서 내 비중 [N%]
            - [키워드 2]: 출현 빈도 [N회], 중요도 [N/10점], 문서 내 비중 [N%]
            {"- 키워드 설명: [기술적 의미 및 적용 사례]" if analysis_depth in ["상세", "심화"] else ""}
            
            ### 2. 비즈니스 키워드
            - [키워드 1]: 출현 빈도 [N회], 중요도 [N/10점], 문서 내 비중 [N%]
            - [키워드 2]: 출현 빈도 [N회], 중요도 [N/10점], 문서 내 비중 [N%]
            {"- 비즈니스 임팩트: [예상 영향도]" if analysis_depth in ["상세", "심화"] else ""}
            
            ### 3. 업계 특화 키워드
            - [키워드 1]: 출현 빈도 [N회], 중요도 [N/10점], 문서 내 비중 [N%]
            - [키워드 2]: 출현 빈도 [N회], 중요도 [N/10점], 문서 내 비중 [N%]
            {f"- {industry} 업종 특성: [상세 설명]" if analysis_depth in ["상세", "심화"] else ""}
            
            ### 4. 트렌드 키워드
            - [키워드 1]: 출현 빈도 [N회], 중요도 [N/10점], 문서 내 비중 [N%]
            - [키워드 2]: 출현 빈도 [N회], 중요도 [N/10점], 문서 내 비중 [N%]
            {"- 최신 트렌드 분석: [산업 동향과의 연관성]" if analysis_depth == "심화" else ""}
            
            **중요도 평가 기준:**
            - 10점: 프로젝트 성공의 핵심 요소
            - 8-9점: 매우 중요한 요구사항
            - 6-7점: 중요한 고려사항
            - 4-5점: 보통 수준의 중요도
            - 1-3점: 참고 수준
            
            {"### 5. 키워드 연관성 분석\n- 주요 키워드 간 관계 맵\n- 키워드 클러스터링 결과" if analysis_depth == "심화" else ""}
            """
        }
    ]
    
    return azure_services.call_openai(messages)


def generate_summary_report(content, industry, analysis_depth, focus_area):
    """요약 보고서 생성"""
    azure_services = st.session_state.azure_services
    return generate_summary_report_with_azure(azure_services, content, industry, analysis_depth, focus_area)

def generate_summary_report_with_azure(azure_services, content, industry, analysis_depth, focus_area):
    """Azure 서비스를 전달받아 요약 보고서 생성"""
    
    messages = [
        {
            "role": "system",
            "content": f"당신은 {industry} 업종의 RFP 분석 보고서 작성 전문가입니다. 종합적인 분석 보고서를 작성해주세요."
        },
        {
            "role": "user",
            "content": f"""
            업종: {industry}
            분석 깊이: {analysis_depth}
            중점 영역: {', '.join(focus_area)}
            
            RFP 문서 내용:
            {content}
            
            다음 구조로 종합 보고서를 작성해주세요:
            
            ## RFP 분석 종합 보고서
            
            ### 1. Executive Summary
            - 프로젝트 개요
            - 주요 요구사항 요약
            - 예상 복잡도 및 리스크
            
            ### 2. 핵심 요구사항 분석
            - Critical 요구사항
            - High Priority 요구사항
            - 기술적 도전 요소
            
            ### 3. 업종별 특화 분석
            - {industry} 업종 특성 반영
            - 규제 요구사항
            - 경쟁사 벤치마킹
            
            ### 4. 제안 전략 가이드
            - 강점 활용 방안
            - 약점 보완 전략
            - 차별화 포인트
            
            ### 5. 다음 단계 권장사항
            - 추가 정보 수집 필요사항
            - 제안서 작성 우선순위
            - 리스크 관리 방안
            """
        }
    ]
    
    return azure_services.call_openai(messages)

def create_keyword_cloud():
    """키워드 클라우드 생성 (샘플)"""
    st.subheader("📊 키워드 빈도 분석")
    
    # 샘플 키워드 데이터
    keywords_data = {
        "키워드": ["디지털전환", "클라우드", "API", "보안", "성능", "통합", "사용자경험", "데이터분석"],
        "빈도": [15, 12, 10, 9, 8, 7, 6, 5],
        "중요도": ["High", "High", "Medium", "High", "Medium", "Medium", "Medium", "Low"]
    }
    
    df = pd.DataFrame(keywords_data)
    st.dataframe(df, use_container_width=True)

def save_to_azure_storage(file_name, file_content, analysis_depth, focus_area, extracted_text=None):
    """Azure Storage에 저장"""
    try:
        azure_services = st.session_state.azure_services
        
        # 디렉토리명 생성 (rfp{timestamp}) - 영어와 숫자만 사용
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        directory_name = f"rfp{timestamp}"
        
        # RFP 파일명을 main_rfp_ 접두사로 변경
        main_rfp_name = f"main_rfp_{file_name}"
        
        # rfp-documents 컨테이너 내에 디렉토리 생성
        container_name = "rfp-documents"
        
        # RFP 파일 업로드 (디버깅 정보 추가)
        st.info(f"📤 Azure Storage에 업로드 중... 파일명: {main_rfp_name}, 크기: {len(file_content)} bytes")
        upload_success = azure_services.upload_file_to_directory(container_name, directory_name, main_rfp_name, file_content)
        
        if upload_success:
            # 추출된 텍스트가 있으면 별도로 저장
            if extracted_text:
                extracted_text_name = f"extracted_text_{file_name}.txt"
                azure_services.upload_file_to_directory(
                    container_name, 
                    directory_name, 
                    extracted_text_name, 
                    extracted_text.encode('utf-8')
                )
            
            # 프로젝트명 한글 요약 메타데이터 생성 (추출된 텍스트 우선 사용)
            content_for_summary = extracted_text if extracted_text else file_content
            project_summary = generate_project_summary(content_for_summary, analysis_depth, focus_area)
            korean_name = generate_korean_project_name(project_summary)
            
            # 메타데이터 저장
            metadata = {
                'korean_name': korean_name,
                'created_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'project_summary': project_summary,
                'original_filename': file_name,
                'analysis_depth': analysis_depth,
                'focus_areas': focus_area
            }
            
            if azure_services.save_directory_metadata_to_path(container_name, directory_name, metadata):
                # 현재 작업 디렉토리를 세션 상태에 저장
                st.session_state.current_directory = directory_name
                st.session_state.current_container = container_name
                st.success(f"📁 RFP가 성공적으로 저장되었습니다: {korean_name} ({directory_name})")
            else:
                st.warning("메타데이터 저장에 실패했습니다.")
        else:
            st.error(f"❌ 파일 저장에 실패했습니다. 파일 크기: {len(file_content)} bytes")
            
    except Exception as e:
        st.error(f"저장 중 오류: {str(e)}")

def generate_project_summary(file_content, analysis_depth, focus_area):
    """프로젝트명 한글 요약 생성"""
    try:
        azure_services = st.session_state.azure_services
        
        # 파일 내용이 바이트인 경우 문자열로 변환
        if isinstance(file_content, bytes):
            try:
                content_text = file_content.decode('utf-8', errors='ignore')
            except:
                content_text = str(file_content)
        else:
            content_text = str(file_content)
        
        messages = [
            {
                "role": "system",
                "content": "당신은 RFP 문서 분석 전문가입니다. RFP 내용을 바탕으로 프로젝트의 핵심을 한글로 요약해주세요."
            },
            {
                "role": "user",
                "content": f"""
                RFP 내용: {content_text[:2000]}...
                분석 깊이: {analysis_depth}
                중점 영역: {', '.join(focus_area)}
                
                다음 형식으로 프로젝트를 한글로 요약해주세요:
                
                ## 프로젝트 요약
                - 프로젝트명: [한글 프로젝트명]
                - 핵심 목표: [프로젝트의 핵심 목표]
                - 주요 기능: [주요 기능 3-5개]
                - 기술 스택: [주요 기술 스택]
                - 예상 규모: [소규모/중규모/대규모]
                """
            }
        ]
        
        return azure_services.call_openai(messages)
    except Exception as e:
        return f"프로젝트 요약 생성 중 오류: {str(e)}"

def generate_enhanced_project_summary(content, industry, analysis_depth, focus_area):
    """재분석용 향상된 프로젝트 요약 생성"""
    try:
        azure_services = st.session_state.azure_services
        
        # 파일 내용이 바이트인 경우 문자열로 변환
        if isinstance(content, bytes):
            try:
                content_text = content.decode('utf-8', errors='ignore')
            except:
                content_text = str(content)
        else:
            content_text = str(content)
        
        messages = [
            {
                "role": "system",
                "content": f"""당신은 {industry} 업종의 RFP 문서 분석 전문가입니다. 
                RFP 내용을 정확히 분석하여 프로젝트의 핵심 정보를 추출하고 한글로 요약해주세요.
                실제 RFP 내용을 바탕으로 구체적이고 정확한 정보를 제공해야 합니다."""
            },
            {
                "role": "user",
                "content": f"""
                업종: {industry}
                RFP 내용: {content_text[:3000]}...
                분석 깊이: {analysis_depth}
                중점 영역: {', '.join(focus_area)}
                
                다음 형식으로 프로젝트를 정확히 한글로 요약해주세요:
                
                ## 프로젝트 요약
                - 프로젝트명: [RFP에서 추출한 실제 프로젝트명]
                - 핵심 목표: [RFP에서 명시된 구체적인 목표]
                - 주요 기능: [RFP에서 요구하는 주요 기능들]
                - 기술 스택: [RFP에서 요구하는 기술 스택]
                - 예상 규모: [RFP 내용을 바탕으로 판단한 규모]
                - 예산 규모: [RFP에서 언급된 예산 정보]
                - 납기일: [RFP에서 명시된 일정]
                - 주요 요구사항: [RFP의 핵심 요구사항 3-5개]
                """
            }
        ]
        
        return azure_services.call_openai(messages)
    except Exception as e:
        return f"향상된 프로젝트 요약 생성 중 오류: {str(e)}"

def generate_korean_project_name(project_summary):
    """프로젝트 한글명 생성"""
    try:
        azure_services = st.session_state.azure_services
        
        messages = [
            {
                "role": "system",
                "content": "당신은 프로젝트명 생성 전문가입니다. 프로젝트 요약을 바탕으로 간결하고 명확한 한글 프로젝트명을 생성해주세요."
            },
            {
                "role": "user",
                "content": f"""
                프로젝트 요약:
                {project_summary}
                
                다음 조건을 만족하는 한글 프로젝트명을 생성해주세요:
                - 10-20자 이내
                - 프로젝트의 핵심을 잘 표현
                - 이해하기 쉬운 용어 사용
                - 금융/IT 업계에 적합한 표현
                
                프로젝트명만 출력해주세요 (설명 없이).
                """
            }
        ]
        
        korean_name = azure_services.call_openai(messages)
        # 불필요한 텍스트 제거
        korean_name = korean_name.strip().replace('프로젝트명:', '').replace('**', '').strip()
        return korean_name[:20]  # 최대 20자로 제한
    except Exception as e:
        return f"프로젝트명 생성 중 오류: {str(e)}"

def save_analysis_results_to_directory(content, industry, analysis_depth, focus_area, requirements, keywords, summary):
    """분석 결과를 디렉토리에 자동 저장"""
    try:
        azure_services = st.session_state.azure_services
        
        # 세션 상태에서 현재 작업 디렉토리 가져오기
        if not hasattr(st.session_state, 'current_directory') or not st.session_state.current_directory:
            # 디렉토리 정보가 없으면 새 디렉토리 생성
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            directory_name = f"rfp{timestamp}"
            container_name = "rfp-documents"
            
            # 세션 상태에 디렉토리 정보 저장
            st.session_state.current_directory = directory_name
            st.session_state.current_container = container_name
            
            st.info(f"📁 새 디렉토리가 생성되었습니다: {directory_name}")
        else:
            # 이미 디렉토리가 설정되어 있으면 그 디렉토리 사용
            container_name = st.session_state.current_container
            directory_name = st.session_state.current_directory
            st.info(f"📁 기존 디렉토리를 사용합니다: {directory_name}")
        
        # 1. 상세 분석 결과 저장
        detailed_content = f"""
# RFP 상세 분석 결과

## 1. 요구사항 추출 결과
{requirements}

## 2. 키워드 분석 결과
{keywords}
        """
        
        # Word 문서 생성 및 저장
        doc = Document()
        doc.add_heading('RFP 상세 분석 결과', 0)
        doc.add_paragraph(detailed_content)
        
        # 임시 파일로 저장
        temp_detailed = "temp_analysis_result_detail.docx"
        doc.save(temp_detailed)
        
        with open(temp_detailed, 'rb') as f:
            detailed_data = f.read()
        
        # 타임스탬프 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Azure에 업로드 (타임스탬프 포함)
        detailed_filename = f"analysis_result_detail_{timestamp}.docx"
        azure_services.upload_file_to_directory(container_name, directory_name, detailed_filename, detailed_data)
        
        # 2. 요약 보고서 저장
        doc_summary = Document()
        doc_summary.add_heading('RFP 분석 요약 보고서', 0)
        doc_summary.add_paragraph(summary)
        
        temp_summary = "temp_analysis_result_summary.docx"
        doc_summary.save(temp_summary)
        
        with open(temp_summary, 'rb') as f:
            summary_data = f.read()
        
        # Azure에 업로드 (타임스탬프 포함)
        summary_filename = f"analysis_result_summary_{timestamp}.docx"
        azure_services.upload_file_to_directory(container_name, directory_name, summary_filename, summary_data)
        
        # 임시 파일 삭제
        import os
        os.remove(temp_detailed)
        os.remove(temp_summary)
        
        st.success("📁 분석 결과가 디렉토리에 자동 저장되었습니다!")
        
    except Exception as e:
        st.error(f"분석 결과 저장 중 오류: {str(e)}")
        st.error("💡 **해결 방법:** 페이지를 새로고침하고 다시 시도해주세요. 문제가 지속되면 관리자에게 문의하세요.")

def create_download_data(filename, content):
    """다운로드용 데이터 생성"""
    try:
        # Word 문서 생성
        doc = Document()
        doc.add_heading('RFP 분석 결과', 0)
        doc.add_paragraph(content)
        
        # 임시 파일로 저장
        temp_filename = f"temp_{filename}"
        doc.save(temp_filename)
        
        # 파일 읽기
        with open(temp_filename, 'rb') as f:
            file_data = f.read()
        
        # 임시 파일 삭제
        import os
        os.remove(temp_filename)
        
        return file_data
        
    except Exception as e:
        st.error(f"다운로드 데이터 생성 중 오류: {str(e)}")
        return b""

def create_download_link(data, filename, label):
    """다운로드 링크 생성 (페이지 리로드 방지)"""
    import base64
    import io
    
    # 데이터를 base64로 인코딩
    b64 = base64.b64encode(data).decode()
    
    # HTML 다운로드 링크 생성
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" style="display: inline-block; padding: 0.5rem 1rem; background-color: #1f77b4; color: white; text-decoration: none; border-radius: 0.25rem; border: none; cursor: pointer;">{label}</a>'
    
    return href

def download_analysis_result(filename, content):
    """분석 결과 다운로드 (기존 방식 유지)"""
    try:
        # Word 문서 생성
        doc = Document()
        doc.add_heading('RFP 분석 결과', 0)
        doc.add_paragraph(content)
        
        # 임시 파일로 저장
        temp_filename = f"temp_{filename}"
        doc.save(temp_filename)
        
        # 파일 읽기
        with open(temp_filename, 'rb') as f:
            file_data = f.read()
        
        # 다운로드 제공
        st.download_button(
            label=f"{filename} 다운로드",
            data=file_data,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # 임시 파일 삭제
        import os
        os.remove(temp_filename)
        
    except Exception as e:
        st.error(f"다운로드 중 오류: {str(e)}")

def extract_text_from_pdf(uploaded_file):
    """PDF 파일에서 텍스트 추출"""
    try:
        # 파일을 바이트로 읽기
        file_bytes = uploaded_file.read()
        return extract_text_from_pdf_bytes(file_bytes)
        
    except Exception as e:
        st.error(f"PDF 텍스트 추출 중 오류: {str(e)}")
        return ""

def extract_text_from_pdf_bytes(file_bytes):
    """PDF 바이트 데이터에서 텍스트 추출"""
    try:
        # pdfplumber를 사용한 텍스트 추출 (더 정확함)
        text_content = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n"
        
        # pdfplumber로 텍스트를 추출하지 못한 경우 PyPDF2 사용
        if not text_content.strip():
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
        
        return text_content.strip()
        
    except Exception as e:
        st.error(f"PDF 바이트 데이터 텍스트 추출 중 오류: {str(e)}")
        return ""

def extract_text_from_docx(uploaded_file):
    """DOCX 파일에서 텍스트 추출"""
    try:
        # 파일을 바이트로 읽기
        file_bytes = uploaded_file.read()
        return extract_text_from_docx_bytes(file_bytes)
        
    except Exception as e:
        st.error(f"DOCX 텍스트 추출 중 오류: {str(e)}")
        return ""

def extract_text_from_docx_bytes(file_bytes):
    """DOCX 바이트 데이터에서 텍스트 추출"""
    try:
        # BytesIO를 사용하여 Document 객체 생성
        doc = Document(io.BytesIO(file_bytes))
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        return text_content.strip()
        
    except Exception as e:
        st.error(f"DOCX 바이트 데이터 텍스트 추출 중 오류: {str(e)}")
        return ""

def extract_text_from_txt(uploaded_file):
    """TXT 파일에서 텍스트 추출"""
    try:
        # 파일을 바이트로 읽기
        file_bytes = uploaded_file.read()
        
        # 파일을 문자열로 디코딩
        if uploaded_file.encoding:
            text_content = str(file_bytes, uploaded_file.encoding)
        else:
            # 인코딩을 모르는 경우 UTF-8로 시도
            text_content = str(file_bytes, 'utf-8')
        
        return text_content.strip()
        
    except Exception as e:
        st.error(f"TXT 텍스트 추출 중 오류: {str(e)}")
        return ""

def extract_text_from_uploaded_file(uploaded_file):
    """업로드된 파일에서 텍스트 추출 (파일 형식에 따라 자동 선택)"""
    file_extension = uploaded_file.name.lower().split('.')[-1]
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(uploaded_file)
    elif file_extension == 'docx':
        return extract_text_from_docx(uploaded_file)
    elif file_extension == 'txt':
        return extract_text_from_txt(uploaded_file)
    else:
        st.error(f"지원하지 않는 파일 형식: {file_extension}")
        return ""
