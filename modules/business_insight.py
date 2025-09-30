"""
비즈니스 인사이트 향상 페이지
"""
import streamlit as st
import json
import pandas as pd
import time
import io
from datetime import datetime
from docx import Document
import PyPDF2
import pdfplumber

def show():
    """비즈니스 인사이트 향상 페이지 표시"""
    st.title("📈 비즈니스 인사이트 향상")
    st.markdown("업계 트렌드와 차별화 전략을 제공합니다.")
    
    # 저장된 RFP 기반 분석만 수행
    show_stored_rfp_analysis()

def show_stored_rfp_analysis():
    """저장된 RFP 기반 분석"""
    st.subheader("📁 저장된 RFP 기반 분석")
    
    try:
        azure_services = st.session_state.azure_services
        directories = azure_services.get_directories()
        
        if not directories:
            st.warning("저장된 RFP가 없습니다.")
            return
        
        # 디렉토리 목록을 테이블로 표시
        st.subheader("📋 저장된 RFP 목록")
        
        # 페이지네이션 설정
        items_per_page = 10
        total_pages = (len(directories) + items_per_page - 1) // items_per_page
        
        if total_pages > 1:
            page = st.selectbox("페이지 선택", range(1, total_pages + 1), key="business_page")
        else:
            page = 1
        
        start_idx = (page - 1) * items_per_page
        end_idx = start_idx + items_per_page
        current_directories = directories[start_idx:end_idx]
        
        # 테이블 데이터 생성 (선택 컬럼 포함)
        table_data = []
        for i, directory in enumerate(current_directories):
            table_data.append({
                "선택": f"RFP {i+1}",
                "한글명": directory['korean_name'],
                "생성일": directory['created_date'],
                "프로젝트 요약": directory['project_summary'][:50] + "..." if len(directory['project_summary']) > 50 else directory['project_summary']
            })
        
        # 테이블과 라디오 버튼을 함께 표시
        st.markdown("**📋 저장된 RFP 목록에서 분석할 RFP를 선택하세요:**")
        
        # 컨테이너로 테이블과 라디오 버튼을 그룹화
        with st.container():
            # 테이블 표시
            df = pd.DataFrame(table_data)
            st.dataframe(df, width='stretch')
            
            # 라디오 버튼을 테이블 바로 아래에 배치
            selected_index = st.radio(
                "분석할 RFP 선택:",
                range(len(current_directories)),
                format_func=lambda x: f"📈 {current_directories[x]['korean_name']} ({current_directories[x]['created_date']})",
                key="business_rfp_selection"
            )
        
        if selected_index is not None:
            selected_directory = current_directories[selected_index]
            
            # 현재 작업 디렉토리를 세션 상태에 저장
            st.session_state.current_directory = selected_directory['name']
            st.session_state.current_container = "rfp-documents"
            
            # RFP 정보 추출
            rfp_info = extract_rfp_info_from_directory(selected_directory)
            
            if st.button("비즈니스 인사이트 생성", type="primary"):
                with st.spinner("비즈니스 인사이트를 생성하고 있습니다..."):
                    generate_business_insights(selected_directory['name'], rfp_info)
                
                st.success("비즈니스 인사이트 생성이 완료되었습니다!")
                
    except Exception as e:
        st.error(f"오류가 발생했습니다: {str(e)}")


def extract_rfp_info_from_directory(directory_info):
    """디렉토리 정보에서 RFP 정보 추출"""
    try:
        azure_services = st.session_state.azure_services
        metadata = azure_services.get_directory_metadata(directory_info['name'])
        
        return {
            "name": directory_info['korean_name'],
            "directory_name": directory_info['name'],
            "date": directory_info['created_date'],
            "industry": "은행",  # 기본값, 실제로는 메타데이터에서 추출
            "project_summary": directory_info['project_summary'],
            "metadata": metadata
        }
    except Exception as e:
        return {
            "name": directory_info['korean_name'],
            "directory_name": directory_info['name'],
            "date": directory_info['created_date'],
            "industry": "은행",
            "project_summary": directory_info['project_summary']
        }

def generate_business_insights(directory_name, rfp_info):
    """비즈니스 인사이트 생성"""
    try:
        azure_services = st.session_state.azure_services
        
        # 저장된 RFP 기반 분석 - analysis_result_summary.docx 파일 참고
        st.info("📄 RFP 분석 결과 파일을 찾고 있습니다...")
        summary_content = get_analysis_summary(directory_name)
        
        if summary_content:
            rfp_info['analysis_summary'] = summary_content
            st.success(f"✅ 분석 파일 발견: {summary_content.get('file_name', 'N/A')}")
            st.info(f"📊 분석 내용 길이: {len(summary_content.get('content', ''))} 자")
        else:
            st.warning("⚠️ 분석 결과 파일을 찾을 수 없습니다. 기본 정보만으로 인사이트를 생성합니다.")
        
        # 항상 새로운 분석 수행
        
        # 탭으로 결과 표시
        tab1, tab2, tab3 = st.tabs([
            "최신 업계 트렌드 요약", 
            "차별화 전략 제안", 
            "자동 생성 스토리라인"
        ])
        
        # 병렬 처리를 위한 placeholder 생성
        with tab1:
            trend_placeholder = st.empty()
            trend_placeholder.info("🔄 업계 트렌드 분석 중...")
        
        with tab2:
            strategy_placeholder = st.empty()
            strategy_placeholder.info("🔄 차별화 전략 분석 중...")
        
        with tab3:
            story_placeholder = st.empty()
            story_placeholder.info("🔄 스토리라인 생성 중...")
        
        # Azure 서비스를 미리 가져와서 병렬 처리에서 사용할 수 있도록 준비
        azure_services = st.session_state.azure_services
        
        # 최적화된 병렬 처리 실행
        from modules.performance import parallel_analysis_executor
        
        analyses = [
            {
                'name': 'industry_trends',
                'func': generate_industry_trends_with_azure,
                'args': (azure_services, rfp_info),
                'kwargs': {}
            },
            {
                'name': 'differentiation_strategy',
                'func': generate_differentiation_strategy_with_azure,
                'args': (azure_services, rfp_info),
                'kwargs': {}
            },
            {
                'name': 'storyline',
                'func': generate_storyline_with_azure,
                'args': (azure_services, rfp_info),
                'kwargs': {}
            }
        ]
        
        def run_analysis_with_azure(azure_services, rfp_info):
            """최적화된 Azure 서비스 분석 실행"""
            results = parallel_analysis_executor(analyses, max_workers=3)
            
            # 결과를 순차적으로 표시 (완료되는 대로)
            # 업계 트렌드 결과 표시
            industry_trends = results['industry_trends']
            trend_placeholder.markdown(industry_trends)
            
            # 차별화 전략 결과 표시
            differentiation_strategy = results['differentiation_strategy']
            strategy_placeholder.markdown(differentiation_strategy)
            
            # 스토리라인 결과 표시
            storyline = results['storyline']
            story_placeholder.markdown(storyline)
            
            return results
        
        # 분석 실행
        results = run_analysis_with_azure(azure_services, rfp_info)
        industry_trends = results['industry_trends']
        differentiation_strategy = results['differentiation_strategy']
        storyline = results['storyline']
        
        # 통합 비즈니스 인사이트 결과 생성
        combined_content = f"""
# 비즈니스 인사이트 결과

## 1. 업계 트렌드 요약
{industry_trends}

## 2. 차별화 전략 제안
{differentiation_strategy}

## 3. 스토리라인
{storyline}
        """
        
        # 비즈니스 인사이트 결과를 디렉토리에 자동 저장
        unique_filename = save_business_insights_to_directory(directory_name, industry_trends, differentiation_strategy, storyline)
        
        # 세션 상태 저장 제거 - 항상 새로운 분석 수행
        
        # 다운로드 버튼을 탭 밖으로 이동
        st.subheader("📥 비즈니스 인사이트 결과 다운로드")
        
        # 다운로드 데이터 미리 생성
        new_insight_download_data = create_insight_download_data(unique_filename, combined_content)
        
        # HTML 다운로드 링크 사용 (페이지 리로드 방지)
        new_insight_link = create_insight_download_link(new_insight_download_data, unique_filename, "📈 비즈니스 인사이트 결과 다운로드")
        st.markdown(new_insight_link, unsafe_allow_html=True)
                
    except Exception as e:
        st.error(f"인사이트 생성 중 오류: {str(e)}")

def get_analysis_summary(directory_name):
    """최신 analysis_result_summary 파일에서 분석 요약 가져오기"""
    try:
        azure_services = st.session_state.azure_services
        container_name = "rfp-documents"
        files = azure_services.list_files_in_directory(container_name, directory_name)
        
        # analysis_result_summary로 시작하는 파일들 중 가장 최신 파일 찾기
        summary_files = [f for f in files if f.startswith('analysis_result_summary') and f.endswith('.docx')]
        if summary_files:
            # 파일명에서 타임스탬프 추출하여 정렬 (최신순)
            summary_files.sort(reverse=True)
            latest_file = summary_files[0]
            
            # DOCX 파일 내용 추출
            file_content = azure_services.download_file_from_directory(container_name, directory_name, latest_file)
            if file_content:
                # DOCX 파일 파싱
                extracted_text = extract_text_from_docx_bytes(file_content)
                if extracted_text:
                    return {
                        'file_name': latest_file,
                        'content': extracted_text,
                        'file_url': f"{container_name}/{directory_name}/{latest_file}"
                    }
        
        return None
    except Exception as e:
        print(f"분석 요약 파일 조회 오류: {e}")
        return None

def extract_text_from_docx_bytes(file_bytes):
    """바이트 데이터에서 DOCX 텍스트 추출"""
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # 표 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n\n".join(full_text)
    except Exception as e:
        print(f"DOCX 텍스트 추출 오류: {e}")
        return None

def generate_industry_trends(rfp_info):
    """최신 업계 트렌드 요약 생성"""
    azure_services = st.session_state.azure_services
    return generate_industry_trends_with_azure(azure_services, rfp_info)

def generate_industry_trends_with_azure(azure_services, rfp_info):
    """Azure 서비스를 전달받아 업계 트렌드 요약 생성"""
    
    # 분석 결과 파일 정보 추출
    analysis_summary = rfp_info.get('analysis_summary', {})
    analysis_content = ""
    file_info = ""
    
    if analysis_summary and isinstance(analysis_summary, dict):
        analysis_content = analysis_summary.get('content', '')
        file_info = f"\n\n📄 **분석 자료**: {analysis_summary.get('file_name', 'N/A')}\n📁 **위치**: {analysis_summary.get('file_url', 'N/A')}"
    
    messages = [
        {
            "role": "system",
            "content": f"""당신은 {rfp_info.get('industry', '금융')} 업종의 비즈니스 트렌드 전문가입니다.
            RFP 분석 결과를 깊이 있게 분석하여, 해당 프로젝트에 맞는 최신 업계 동향과 트렌드를 구체적으로 요약해주세요.
            
            분석 시 다음을 중점적으로 고려하세요:
            - RFP에 명시된 요구사항과 제약사항
            - 프로젝트의 목표와 범위
            - 고객사의 비즈니스 특성
            - 해당 업종의 최신 기술 동향"""
        },
        {
            "role": "user",
            "content": f"""
            ## RFP 기본 정보
            - **프로젝트명**: {rfp_info.get('name', 'N/A')}
            - **업종**: {rfp_info.get('industry', '금융')}
            - **생성일**: {rfp_info.get('date', 'N/A')}
            {file_info}
            
            ## RFP 상세 분석 결과
            {analysis_content if analysis_content else '분석 결과 파일을 찾을 수 없습니다. 기본 정보만 활용합니다.'}
            
            ---
            
            위 RFP 분석 결과를 바탕으로, 다음 항목에 대해 **구체적이고 실질적인** 업계 트렌드를 분석해주세요:
            
            ## 1. 기술 트렌드
            - RFP 요구사항과 관련된 핵심 기술 동향
            - 프로젝트에 적용 가능한 신기술
            - 기술 도입의 우선순위 및 투자 방향
            
            ## 2. 시장 트렌드
            - 해당 업종/분야의 시장 규모 및 성장률
            - 경쟁사 및 선도 기업의 유사 프로젝트 사례
            - 시장 기회와 위험 요소
            
            ## 3. 규제 및 정책 트렌드
            - RFP와 관련된 최신 규제 및 컴플라이언스 요구사항
            - 정책 변화가 프로젝트에 미치는 영향
            - 준수해야 할 표준 및 인증
            
            ## 4. 고객 행동 및 니즈 변화
            - 고객사 업종의 사용자 니즈 변화
            - 디지털 전환 가속화에 따른 요구사항
            - 사용자 경험(UX) 및 접근성 트렌드
            
            **중요**: 일반적인 내용이 아닌, 이 RFP 분석 결과에 기반한 구체적이고 맥락에 맞는 인사이트를 제공해주세요.
            """
        }
    ]
    
    return azure_services.call_openai(messages)

def generate_differentiation_strategy(rfp_info):
    """차별화 전략 제안 생성"""
    azure_services = st.session_state.azure_services
    return generate_differentiation_strategy_with_azure(azure_services, rfp_info)

def generate_differentiation_strategy_with_azure(azure_services, rfp_info):
    """Azure 서비스를 전달받아 차별화 전략 제안 생성"""
    
    # 분석 결과 파일 정보 추출
    analysis_summary = rfp_info.get('analysis_summary', {})
    analysis_content = ""
    file_info = ""
    
    if analysis_summary and isinstance(analysis_summary, dict):
        analysis_content = analysis_summary.get('content', '')
        file_info = f"\n\n📄 **분석 자료**: {analysis_summary.get('file_name', 'N/A')}\n📁 **위치**: {analysis_summary.get('file_url', 'N/A')}"
    
    messages = [
        {
            "role": "system",
            "content": f"""당신은 {rfp_info.get('industry', '금융')} 업종의 경쟁 전략 전문가입니다.
            RFP 분석 결과를 바탕으로 고객의 요구사항과 제약사항을 깊이 이해하고,
            경쟁사 대비 명확한 차별화 전략을 구체적으로 제안해주세요.
            
            제안 시 다음을 고려하세요:
            - RFP에 명시된 필수 요구사항 및 우대사항
            - 고객사의 현재 시스템 및 프로세스
            - 경쟁 환경 및 시장 동향
            - 실현 가능한 구체적 전략"""
        },
        {
            "role": "user",
            "content": f"""
            ## RFP 기본 정보
            - **프로젝트명**: {rfp_info.get('name', 'N/A')}
            - **업종**: {rfp_info.get('industry', '금융')}
            - **생성일**: {rfp_info.get('date', 'N/A')}
            {file_info}
            
            ## RFP 상세 분석 결과
            {analysis_content if analysis_content else '분석 결과 파일을 찾을 수 없습니다. 기본 정보만 활용합니다.'}
            
            ---
            
            위 RFP 분석 결과를 바탕으로, 다음 관점에서 **실현 가능하고 구체적인** 차별화 전략을 제안해주세요:
            
            ## 1. 기술적 차별화
            - RFP 요구사항을 초과 달성하는 혁신 기술
            - 성능/안정성/확장성 우위 요소
            - 레퍼런스 및 검증된 기술력
            
            ## 2. 서비스 차별화
            - 고객사 특성에 맞춤화된 서비스 모델
            - 사용자 경험(UX) 개선 방안
            - 운영 및 유지보수 효율성 향상
            
            ## 3. 비즈니스 모델 차별화
            - 비용 효율적인 가격 전략
            - 단계적 구축을 통한 리스크 최소화
            - 성과 기반 계약 모델
            
            ## 4. 실행 계획 및 리스크 관리
            - 구체적인 단계별 구현 방안
            - 주요 리스크 요소 및 대응 전략
            - KPI 및 성과 측정 지표
            
            **중요**: 일반론이 아닌, 이 RFP의 특정 요구사항과 제약사항에 기반한 맞춤형 전략을 제시해주세요.
            """
        }
    ]
    
    return azure_services.call_openai(messages)

def generate_storyline(rfp_info):
    """자동 생성 스토리라인"""
    azure_services = st.session_state.azure_services
    return generate_storyline_with_azure(azure_services, rfp_info)

def generate_storyline_with_azure(azure_services, rfp_info):
    """Azure 서비스를 전달받아 스토리라인 생성"""
    
    # 분석 결과 파일 정보 추출
    analysis_summary = rfp_info.get('analysis_summary', {})
    analysis_content = ""
    file_info = ""
    
    if analysis_summary and isinstance(analysis_summary, dict):
        analysis_content = analysis_summary.get('content', '')
        file_info = f"\n\n📄 **분석 자료**: {analysis_summary.get('file_name', 'N/A')}\n📁 **위치**: {analysis_summary.get('file_url', 'N/A')}"
    
    messages = [
        {
            "role": "system",
            "content": f"""당신은 {rfp_info.get('industry', '금융')} 업종의 제안서 스토리텔링 전문가입니다.
            RFP 분석 결과를 깊이 이해하고, 고객의 Pain Point부터 우리 솔루션의 가치까지 
            논리적이고 설득력 있는 제안 스토리라인을 작성해주세요.
            
            스토리라인 작성 원칙:
            - RFP에 명시된 배경과 목적을 정확히 반영
            - 고객사의 현재 상황과 문제점을 구체적으로 파악
            - 우리 솔루션이 왜 최적의 선택인지 논리적으로 연결
            - 실행 가능성과 성공 가능성을 명확히 제시"""
        },
        {
            "role": "user",
            "content": f"""
            ## RFP 기본 정보
            - **프로젝트명**: {rfp_info.get('name', 'N/A')}
            - **업종**: {rfp_info.get('industry', '금융')}
            - **생성일**: {rfp_info.get('date', 'N/A')}
            {file_info}
            
            ## RFP 상세 분석 결과
            {analysis_content if analysis_content else '분석 결과 파일을 찾을 수 없습니다. 기본 정보만 활용합니다.'}
            
            ---
            
            위 RFP 분석 결과를 바탕으로, 다음 구조로 **설득력 있는 제안 스토리라인**을 작성해주세요:
            
            ## 1. 현재 상황 분석 (Pain Point)
            - RFP에서 파악된 고객사의 현재 문제점과 과제
            - 기존 시스템/프로세스의 한계와 개선 필요성
            - 해결하지 못할 경우의 비즈니스 리스크
            
            ## 2. 프로젝트 비전 (Vision)
            - RFP가 목표로 하는 이상적인 미래 상태
            - 달성하고자 하는 비즈니스 목표와 성과
            - 디지털 전환 및 혁신의 방향성
            
            ## 3. 우리의 솔루션 (Solution)
            - RFP 요구사항을 충족하는 제안 개요
            - 핵심 기능과 차별화된 기술적 특징
            - 고객사 환경에 최적화된 구현 방안
            
            ## 4. 제공 가치 (Value Proposition)
            - 고객사에게 제공하는 구체적 가치
            - 정량적 효과 (비용 절감, 효율 향상 등)
            - 정성적 효과 (사용자 만족도, 경쟁력 강화 등)
            
            ## 5. 실행 계획 (Execution & Success)
            - 단계별 구현 로드맵과 주요 마일스톤
            - 리스크 관리 및 품질 보증 방안
            - 프로젝트 성공을 보장하는 우리의 역량과 경험
            
            ## 6. 결론 (Call to Action)
            - 왜 우리를 선택해야 하는가에 대한 명확한 메시지
            - 파트너로서의 장기적 가치 제안
            
            **중요**: RFP 내용에 기반한 구체적이고 맥락에 맞는 스토리를 전개하되, 
            고객을 설득할 수 있는 감성적 공감과 논리적 근거를 균형있게 제시해주세요.
            """
        }
    ]
    
    return azure_services.call_openai(messages)

def save_business_insights_to_directory(directory_name, industry_trends, differentiation_strategy, storyline):
    """비즈니스 인사이트 결과를 별도 디렉토리에 자동 저장"""
    try:
        azure_services = st.session_state.azure_services
        
        # 세션 상태에서 현재 작업 디렉토리 가져오기
        if not hasattr(st.session_state, 'current_directory') or not st.session_state.current_directory:
            # 디렉토리 정보가 없으면 저장을 건너뛰고 조용히 종료
            return
        
        # 기존 RFP 분석 디렉토리 사용 (같은 디렉토리 유지)
        container_name = st.session_state.current_container
        directory_name = st.session_state.current_directory
        
        # 타임스탬프 생성
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # 통합 비즈니스 인사이트 결과 생성
        combined_content = f"""
# 비즈니스 인사이트 결과

## 1. 업계 트렌드 요약
{industry_trends}

## 2. 차별화 전략 제안
{differentiation_strategy}

## 3. 스토리라인
{storyline}
        """
        
        # Word 문서 생성 및 저장
        doc = Document()
        doc.add_heading('비즈니스 인사이트 결과', 0)
        doc.add_paragraph(combined_content)
        
        # 임시 파일로 저장
        temp_filename = "temp_business_insight_result.docx"
        doc.save(temp_filename)
        
        with open(temp_filename, 'rb') as f:
            file_data = f.read()
        
        # Azure에 업로드 (타임스탬프 추가하여 고유 파일명 생성)
        unique_filename = f"business_insight_result_{timestamp}.docx"
        azure_services.upload_file_to_directory(container_name, directory_name, unique_filename, file_data)
        
        # 임시 파일 삭제
        import os
        os.remove(temp_filename)
        
        st.success(f"📁 비즈니스 인사이트 결과가 디렉토리에 저장되었습니다: {unique_filename}")
        return unique_filename
        
    except Exception as e:
        st.error(f"비즈니스 인사이트 저장 중 오류: {str(e)}")
        return None

def create_insight_download_data(filename, content):
    """비즈니스 인사이트 다운로드용 데이터 생성"""
    try:
        # Word 문서 생성
        doc = Document()
        doc.add_heading('비즈니스 인사이트', 0)
        doc.add_paragraph(content)
        
        # 임시 파일로 저장
        temp_filename = f"temp_{filename}"
        doc.save(temp_filename)
        
        # 파일 읽기
        with open(temp_filename, 'rb') as f:
            file_data = f.read()
        
        # 임시 파일 삭제
        import os
        os.remove(temp_filename)
        
        return file_data
        
    except Exception as e:
        st.error(f"다운로드 데이터 생성 중 오류: {str(e)}")
        return b""

def create_insight_download_link(data, filename, label):
    """비즈니스 인사이트 다운로드 링크 생성 (페이지 리로드 방지)"""
    import base64
    
    # 데이터를 base64로 인코딩
    b64 = base64.b64encode(data).decode()
    
    # HTML 다운로드 링크 생성
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}" style="display: inline-block; padding: 0.5rem 1rem; background-color: #1f77b4; color: white; text-decoration: none; border-radius: 0.25rem; border: none; cursor: pointer;">{label}</a>'
    
    return href

def extract_text_from_pdf_bytes(file_bytes):
    """PDF 바이트 데이터에서 텍스트 추출"""
    try:
        # pdfplumber를 사용한 텍스트 추출 (더 정확함)
        text_content = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n"
        
        # pdfplumber로 텍스트를 추출하지 못한 경우 PyPDF2 사용
        if not text_content.strip():
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
        
        return text_content.strip()
        
    except Exception as e:
        st.error(f"PDF 바이트 데이터 텍스트 추출 중 오류: {str(e)}")
        return ""

def extract_text_from_docx_bytes(file_bytes):
    """DOCX 바이트 데이터에서 텍스트 추출"""
    try:
        # BytesIO를 사용하여 Document 객체 생성
        doc = Document(io.BytesIO(file_bytes))
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        return text_content.strip()
        
    except Exception as e:
        st.error(f"DOCX 바이트 데이터 텍스트 추출 중 오류: {str(e)}")
        return ""

def download_insight(filename, content):
    """인사이트 결과 다운로드 (기존 방식 유지)"""
    try:
        # Word 문서 생성
        doc = Document()
        doc.add_heading('비즈니스 인사이트', 0)
        doc.add_paragraph(content)
        
        # 임시 파일로 저장
        temp_filename = f"temp_{filename}"
        doc.save(temp_filename)
        
        # 파일 읽기
        with open(temp_filename, 'rb') as f:
            file_data = f.read()
        
        # 다운로드 제공
        st.download_button(
            label=f"{filename} 다운로드",
            data=file_data,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # 임시 파일 삭제
        import os
        os.remove(temp_filename)
        
    except Exception as e:
        st.error(f"다운로드 중 오류: {str(e)}")





